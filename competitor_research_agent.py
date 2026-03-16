import anthropic
import os
import json
import sys
import requests
from datetime import datetime
from dotenv import load_dotenv
from ddgs import DDGS
from bs4 import BeautifulSoup
from docx import Document
from docx.shared import Pt, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH

load_dotenv(os.path.join(os.path.dirname(__file__), '.env'))

client = anthropic.Anthropic(api_key=os.getenv("ANTHROPIC_API_KEY"))


# --- TOOLS ---

def search_web(query, notes):
    print(f"    [search_web] Query: {query}")
    try:
        with DDGS() as ddgs:
            results = list(ddgs.text(query, max_results=5))
        return json.dumps([{"title": r["title"], "url": r["href"], "snippet": r["body"]} for r in results])
    except Exception as e:
        return f"Search failed: {e}"

def extract_page_content(url, notes):
    print(f"    [extract_page] URL: {url}")
    try:
        response = requests.get(url, timeout=10, headers={"User-Agent": "Mozilla/5.0"})
        soup = BeautifulSoup(response.text, "html.parser")
        text = " ".join(soup.get_text().split())[:2000]
        return text
    except Exception as e:
        return f"Could not extract page: {e}"

def save_note(competitor_name, competitor_type, aspects_data, notes):
    print(f"    [save_note] Saving note for: {competitor_name}")
    notes[competitor_name] = {
        "type": competitor_type,
        "aspects": aspects_data
    }
    return f"Note saved for '{competitor_name}'"

def propose_competitor_list(competitors, notes, approved_competitors):
    competitor_list = list(competitors) if isinstance(competitors, list) else [competitors]

    while True:
        print(f"\n{'='*60}")
        print(f"  [APPROVAL GATE 1] Agent identified these competitors:\n")
        for i, c in enumerate(competitor_list, 1):
            print(f"    {i}. {c}")
        print(f"\n  Commands: press Enter to approve | 'add [name]' | 'remove [name]'")
        print(f"{'='*60}")

        user_input = input("\n  Your decision: ").strip().lower()

        if not user_input:
            approved_competitors.extend(competitor_list)
            print(f"\n  [gate 1] Approved. Proceeding to research all {len(competitor_list)} competitors.\n")
            return f"Competitor list approved: {json.dumps(competitor_list)}. Research ALL of them — do not stop early."
        elif user_input.startswith("add "):
            name = user_input[4:].strip().strip("[]()").title()
            competitor_list.append(name)
            print(f"  Added: {name}")
        elif user_input.startswith("remove "):
            name = user_input[7:].strip().lower()
            before = len(competitor_list)
            competitor_list = [c for c in competitor_list if c.lower() != name]
            if len(competitor_list) < before:
                print(f"  Removed: {name.title()}")
            else:
                print(f"  Not found: {name.title()}")
        else:
            print("  Type 'add [name]', 'remove [name]', or press Enter to approve.")


# --- TOOL DEFINITIONS FOR CLAUDE ---

tools = [
    {
        "name": "search_web",
        "description": "Search the web using DuckDuckGo. Returns titles, URLs, and snippets. Use specific queries like '[competitor name] [company] competitor features fees target audience comparison'.",
        "input_schema": {
            "type": "object",
            "properties": {
                "query": {
                    "type": "string",
                    "description": "A specific search query including the competitor name and what you want to know."
                }
            },
            "required": ["query"]
        }
    },
    {
        "name": "extract_page_content",
        "description": "Extract the full text content from a web page URL. Prefer Wikipedia, official sites, or comparison articles over social media.",
        "input_schema": {
            "type": "object",
            "properties": {
                "url": {
                    "type": "string",
                    "description": "The full URL of the page to extract content from."
                }
            },
            "required": ["url"]
        }
    },
    {
        "name": "propose_competitor_list",
        "description": "After your first search, call this to propose the competitor list to the user for approval before starting research. The user can add or remove competitors. You MUST call this before researching any individual competitor.",
        "input_schema": {
            "type": "object",
            "properties": {
                "competitors": {
                    "type": "array",
                    "items": {"type": "string"},
                    "description": "The list of competitor names you identified from your initial search."
                }
            },
            "required": ["competitors"]
        }
    },
    {
        "name": "save_note",
        "description": "Save structured findings about a competitor. Call this once per competitor immediately after researching them. Data must be structured for table output.",
        "input_schema": {
            "type": "object",
            "properties": {
                "competitor_name": {
                    "type": "string",
                    "description": "The name of the competitor (e.g. 'Amazon', 'Netflix')."
                },
                "competitor_type": {
                    "type": "string",
                    "description": "A short one-line description of what type of company this is (e.g. 'Broad horizontal e-commerce marketplace', 'Subscription video streaming platform')."
                },
                "aspects_data": {
                    "type": "object",
                    "description": "A JSON object where each key is an aspect name and each value is the detailed findings for that aspect. Every requested aspect must be a key.",
                    "additionalProperties": {"type": "string"}
                }
            },
            "required": ["competitor_name", "competitor_type", "aspects_data"]
        }
    }
]


# --- TOOL EXECUTOR ---

def run_tool(name, inputs, notes, approved_competitors):
    if name == "search_web":
        return search_web(inputs["query"], notes)
    elif name == "extract_page_content":
        return extract_page_content(inputs["url"], notes)
    elif name == "propose_competitor_list":
        return propose_competitor_list(inputs["competitors"], notes, approved_competitors)
    elif name == "save_note":
        return save_note(
            inputs["competitor_name"],
            inputs["competitor_type"],
            inputs["aspects_data"],
            notes
        )
    else:
        return f"Unknown tool: {name}"


# --- QUALITY CHECKER ---

def check_quality(notes, aspects):
    if not notes:
        return 0
    response = client.messages.create(
        model="claude-sonnet-4-6",
        max_tokens=256,
        messages=[{
            "role": "user",
            "content": f"""Score the quality of this competitor research from 0-100.
Each competitor must cover these aspects: {', '.join(aspects)}.
Notes collected so far ({len(notes)} competitors):
{json.dumps(notes, indent=2)}

Rules:
- All key competitors covered with all aspects fully addressed = 100
- Missing competitors or incomplete aspects = deduct points proportionally
- Respond with ONLY a number between 0 and 100. Nothing else."""
        }]
    )
    try:
        return int(response.content[0].text.strip())
    except:
        return 0


# =============================================================================
# CONTEXT WINDOW MANAGEMENT — 3 STRATEGIES
# =============================================================================

# --- STRATEGY 1: SUMMARIZATION ---
# Every 3 iterations, compress conversation history into a summary.
# Prevents the message list from growing too large on long runs.

def summarize_history(messages, company, notes):
    print(f"\n  [context] Summarizing conversation history (iteration checkpoint)...")

    # Build a readable version of history for Claude to summarize
    history_text = []
    for m in messages:
        if isinstance(m["content"], str):
            history_text.append(f"{m['role'].upper()}: {m['content']}")
        elif isinstance(m["content"], list):
            for block in m["content"]:
                if isinstance(block, dict):
                    if block.get("type") == "tool_result":
                        history_text.append(f"TOOL RESULT: {str(block.get('content', ''))[:300]}")
                elif hasattr(block, "type"):
                    if block.type == "text":
                        history_text.append(f"ASSISTANT: {block.text[:300]}")
                    elif block.type == "tool_use":
                        history_text.append(f"TOOL CALL: {block.name}({json.dumps(block.input)[:200]})")

    response = client.messages.create(
        model="claude-sonnet-4-6",
        max_tokens=512,
        messages=[{
            "role": "user",
            "content": f"""Summarize what this research agent has done so far. Be concise.
Include: which competitors were identified, which have been researched, which are still pending, and any key findings.

History:
{chr(10).join(history_text[-30:])}

Respond in 3-5 bullet points."""
        }]
    )

    summary = response.content[0].text
    print(f"  [context] History compressed. Summary:\n  {summary[:300]}...")

    # Replace full history with summary + current notes state
    compressed = [
        {
            "role": "user",
            "content": f"""Previous session summary:
{summary}

Competitors researched so far: {list(notes.keys()) if notes else 'None yet'}

Continue researching remaining competitors for {company}."""
        }
    ]
    return compressed


# --- STRATEGY 2: SELECTIVE MEMORY ---
# Keep only the last N tool results in active context.
# Older tool results are archived to a file, not lost — just moved out of context.

def trim_tool_results(messages, keep_last=3, archive_file="context_archive.jsonl"):
    tool_result_indices = [
        i for i, m in enumerate(messages)
        if isinstance(m.get("content"), list)
        and any(
            isinstance(b, dict) and b.get("type") == "tool_result"
            for b in m["content"]
        )
    ]

    if len(tool_result_indices) <= keep_last:
        return messages  # nothing to trim

    # Archive older tool results
    to_archive = tool_result_indices[:-keep_last]
    with open(archive_file, "a") as f:
        for idx in to_archive:
            f.write(json.dumps(messages[idx]) + "\n")

    # Remove tool_result messages AND their preceding tool_use (assistant) messages.
    # The API requires every tool_use to have a matching tool_result immediately after.
    # Removing tool_results without removing their tool_use blocks causes a 400 error.
    indices_to_remove = set(to_archive)
    for idx in to_archive:
        if idx > 0 and messages[idx - 1].get("role") == "assistant":
            indices_to_remove.add(idx - 1)

    trimmed = [m for i, m in enumerate(messages) if i not in indices_to_remove]

    archived_count = len(to_archive)
    if archived_count > 0:
        print(f"  [context] Archived {archived_count} old tool results. Keeping last {keep_last} in active context.")

    return trimmed


# --- STRATEGY 3: SCRATCHPAD ---
# Agent maintains an explicit short-term memory dict.
# This is injected into every prompt so Claude always knows current state
# without relying on long conversation history.

def build_scratchpad(notes, aspects, company, approved_competitors, stall_warning=False):
    researched = list(notes.keys())
    pending = [c for c in approved_competitors if c not in notes]

    scratchpad = {
        "company_being_researched": company,
        "aspects_required": aspects,
        "approved_competitors": approved_competitors,
        "competitors_researched": researched,
        "competitors_pending": pending,
        "progress": f"{len(researched)} of {len(approved_competitors)} done" if approved_competitors else "awaiting competitor list approval"
    }
    if stall_warning:
        scratchpad["URGENT"] = "No notes saved in 5+ iterations. Stop searching immediately. Call save_note NOW using data already gathered. Do not make any more search or extract calls until a note is saved."
    return json.dumps(scratchpad, indent=2)


# =============================================================================
# BUILD WORD DOC WITH TABLES
# =============================================================================

def add_competitor_table(doc, competitor_name, competitor_type, aspects_data):
    doc.add_heading(competitor_name, level=2)

    p = doc.add_paragraph()
    p.add_run("Type: ").bold = True
    p.add_run(competitor_type)
    doc.add_paragraph("")

    table = doc.add_table(rows=1, cols=2)
    table.style = "Light Grid Accent 1"

    header = table.rows[0].cells
    header[0].text = "Aspect"
    header[1].text = "Details"
    for cell in header:
        for para in cell.paragraphs:
            for run in para.runs:
                run.bold = True

    for aspect, details in aspects_data.items():
        row = table.add_row().cells
        row[0].text = aspect.replace("_", " ").title()
        row[1].text = details

    doc.add_paragraph("")


def add_landscape_table(doc, notes, aspects, company):
    doc.add_heading("Competitive Landscape Overview", level=1)
    doc.add_paragraph(f"Summary comparison of all key competitors to {company}.")
    doc.add_paragraph("")

    col_headers = ["Competitor", "Type"] + [a.replace("_", " ").title() for a in aspects]

    table = doc.add_table(rows=1, cols=len(col_headers))
    table.style = "Light Grid Accent 1"

    header = table.rows[0].cells
    for i, h in enumerate(col_headers):
        header[i].text = h
        for para in header[i].paragraphs:
            for run in para.runs:
                run.bold = True

    for name, data in notes.items():
        row = table.add_row().cells
        row[0].text = name
        row[1].text = data.get("type", "")
        for i, aspect in enumerate(aspects):
            aspect_value = ""
            for key, value in data.get("aspects", {}).items():
                if aspect.lower() in key.lower() or key.lower() in aspect.lower():
                    aspect_value = value[:200] + "..." if len(value) > 200 else value
                    break
            row[2 + i].text = aspect_value

    doc.add_paragraph("")


def show_research_summary_and_confirm(notes, aspects, quality, quality_threshold):
    print(f"\n{'='*60}")
    print(f"  [APPROVAL GATE 2] Research complete. Here's what was found:\n")
    for name, data in notes.items():
        print(f"  {name}")
        print(f"  Type: {data.get('type', '')}")
        for aspect, value in data.get("aspects", {}).items():
            short = value.split(".")[0][:120]
            print(f"    • {aspect}: {short}")
        print()

    if quality >= quality_threshold:
        print(f"  Quality score: {quality}/100  (threshold: {quality_threshold}) — good to go")
    else:
        print(f"  Quality score: {quality}/100  Below threshold ({quality_threshold}) — some aspects may be incomplete")

    print(f"\n{'='*60}")
    print(f"  Options:")
    print(f"    yes        -> generate Word document")
    print(f"    no         -> exit without generating")
    print(f"    [feedback] -> request more research (e.g. 'go deeper on Stripe pricing')")

    decision = input("\n  Your decision: ").strip()
    lower = decision.lower()

    if lower in ("yes", "y"):
        return "generate"
    elif lower in ("no", "n"):
        return "skip"
    else:
        return decision  # feedback string — triggers re-research


def save_report(notes, company, aspects):
    timestamp = datetime.now().strftime("%Y%m%d_%H%M%S")
    filename = f"report_{company.replace(' ', '_')}_{timestamp}.docx"

    doc = Document()

    doc.add_heading(f"Competitor Research: {company}", 0)
    doc.add_paragraph(f"Generated: {datetime.now().strftime('%Y-%m-%d %H:%M:%S')}")
    doc.add_paragraph(f"Research aspects: {', '.join(aspects)}")
    doc.add_paragraph("")

    add_landscape_table(doc, notes, aspects, company)

    doc.add_page_break()
    doc.add_heading("Detailed Competitor Profiles", level=1)

    for name, data in notes.items():
        add_competitor_table(
            doc,
            name,
            data.get("type", ""),
            data.get("aspects", {})
        )

    doc.save(filename)
    print(f"\n  [report] Saved to {filename}")
    return filename


# --- USER INPUT ---

def get_user_inputs():
    print("\n=== COMPETITOR RESEARCH AGENT ===\n")

    company = input("Enter the company to research competitors for: ").strip()
    if not company:
        print("No company entered. Exiting.")
        sys.exit(1)

    default_aspects = [
        "Key Features & Offerings",
        "Target Audience",
        "Fee / Pricing Structure",
        f"Key Differentiator vs {company}"
    ]

    add_aspects = input("\nWould you like to add custom research aspects? (e.g. 'Brand trust', 'International strategy') — type 'yes' or press Enter to skip: ").strip().lower()

    extra_aspects = []
    if add_aspects == "yes":
        print("Enter up to 5 aspects (press Enter on any to stop):\n")
        for i in range(1, 6):
            aspect = input(f"  Aspect {i}: ").strip()
            if not aspect:
                break
            extra_aspects.append(aspect)

    aspects = default_aspects + extra_aspects
    print(f"\nResearch aspects: {aspects}")

    return company, aspects


# --- AGENTIC LOOP ---

def run_agent(company, aspects, max_iterations=40, quality_threshold=95,
              summarize_every=8, keep_tool_results=3):
    notes = {}
    approved_competitors = []
    last_note_count = 0
    stall_count = 0
    stall_warning = False
    aspects_str = "\n".join(f"- {a}" for a in aspects)
    goal = f"Identify the key competitors to {company} and research each one covering these aspects:\n{aspects_str}"

    print(f"\nGoal: {goal}")
    print(f"Max iterations: {max_iterations} | Quality scoring: enabled (informational only)")
    print(f"Context: summarize every {summarize_every} iterations | keep last {keep_tool_results} tool results\n")

    messages = [{"role": "user", "content": goal}]
    iteration = 0

    while iteration < max_iterations:
        iteration += 1
        print(f"\n--- Iteration {iteration} of {max_iterations} ---")

        # STRATEGY 1: Summarize every N iterations
        if iteration > 1 and iteration % summarize_every == 0:
            messages = summarize_history(messages, company, notes)

        # STRATEGY 2: Trim old tool results, keep only last N
        messages = trim_tool_results(messages, keep_last=keep_tool_results)

        # Quality check when new note saved
        if len(notes) > last_note_count:
            last_note_count = len(notes)
            stall_count = 0
            stall_warning = False
            quality = check_quality(notes, aspects)
            print(f"\n  [quality check] Notes saved: {len(notes)} of {len(approved_competitors)} competitors | Score: {quality}/100")

            # Stop only when ALL approved competitors are researched
            if approved_competitors and len(notes) >= len(approved_competitors):
                print(f"  [complete] All {len(approved_competitors)} approved competitors researched. Writing report.")
                messages.append({
                    "role": "user",
                    "content": "All approved competitors have been researched and saved. Output a brief final summary and stop."
                })
        else:
            # Only count stall AFTER Gate 1 is approved — pre-Gate-1 iterations are normal setup
            if approved_competitors:
                stall_count += 1
                if stall_count >= 5:
                    stall_warning = True
                    stall_count = 0
                    print(f"\n  [stall detected] No notes saved in 5 iterations. Injecting warning into scratchpad.")

        # STRATEGY 3: Inject scratchpad into every Claude call
        scratchpad = build_scratchpad(notes, aspects, company, approved_competitors, stall_warning)

        response = client.messages.create(
            model="claude-sonnet-4-6",
            max_tokens=4096,
            system=f"""You are a market research agent.

Your goal is to identify and research the key competitors to {company}.

CURRENT STATE (scratchpad — always read this first):
{scratchpad}

RESEARCH PROCESS — FOLLOW EXACTLY IN ORDER:
1. Do ONE broad search to identify the key competitors to {company}
2. Call propose_competitor_list with the competitors you found — WAIT for user approval before proceeding
3. Pick ONE competitor. Search for it. Extract ONE page. Call save_note immediately.
4. Pick the NEXT competitor. Repeat step 3.
5. Continue until all approved competitors are saved. Then output a brief summary and stop.

CRITICAL RULES:
- You MUST call propose_competitor_list after step 1 before researching any individual competitor
- Research and save ONE competitor at a time — never search for multiple competitors in the same turn
- After extracting one page, call save_note IMMEDIATELY — do not read more pages first
- The only allowed pattern is: search → extract → save_note → next competitor
- Do NOT search for a new competitor until the current one has a saved note
- Research ALL competitors in the approved list — never stop early
- Check competitors_pending in the scratchpad to see who still needs to be researched
- Maximum {max_iterations} iterations total
- When saving notes, populate aspects_data with ALL of these as keys:
{aspects_str}
- Make each aspect value detailed and specific — these go directly into a table in the report""",
            tools=tools,
            messages=messages
        )

        if response.stop_reason == "end_turn":
            print(f"\n=== AGENT FINISHED in {iteration} iterations ===")
            for block in response.content:
                if hasattr(block, "text"):
                    print(f"\n{block.text}")
            break

        if response.stop_reason == "tool_use":
            messages.append({"role": "assistant", "content": response.content})

            tool_results = []
            for block in response.content:
                if block.type == "tool_use":
                    result = run_tool(block.name, block.input, notes, approved_competitors)
                    tool_results.append({
                        "type": "tool_result",
                        "tool_use_id": block.id,
                        "content": str(result)
                    })

            messages.append({"role": "user", "content": tool_results})

    if iteration >= max_iterations:
        print(f"\n=== AGENT STOPPED: hit max {max_iterations} iterations ===")

    if notes:
        final_quality = check_quality(notes, aspects)

        while True:
            decision = show_research_summary_and_confirm(notes, aspects, final_quality, quality_threshold)

            if decision == "generate":
                save_report(notes, company, aspects)
                break

            elif decision == "skip":
                print("\n  [gate 2] Word document skipped. Research notes preserved.")
                break

            else:
                # User gave feedback — run more research beyond original limit
                re_research_limit = 10
                print(f"\n  [gate 2] Resuming research ({re_research_limit} additional iterations)...")
                print(f"  Feedback: {decision}\n")

                messages.append({
                    "role": "user",
                    "content": f"User reviewed the research and gave this feedback: '{decision}'. Do additional research to address it. Fill gaps, deepen shallow findings, or re-research any incomplete aspects. You have {re_research_limit} more iterations."
                })

                for extra_iter in range(re_research_limit):
                    print(f"\n--- Re-research iteration {extra_iter + 1} of {re_research_limit} ---")

                    messages = trim_tool_results(messages, keep_last=keep_tool_results)
                    scratchpad = build_scratchpad(notes, aspects, company, approved_competitors, stall_warning=False)

                    response = client.messages.create(
                        model="claude-sonnet-4-6",
                        max_tokens=4096,
                        system=f"""You are a market research agent doing additional research based on user feedback.

CURRENT STATE (scratchpad):
{scratchpad}

User feedback: {decision}

Address the feedback. You may search for more detail, extract additional pages, or update existing notes using save_note (overwriting is fine).
When done, output a brief summary and stop.

Aspects required:
{aspects_str}""",
                        tools=tools,
                        messages=messages
                    )

                    if response.stop_reason == "end_turn":
                        for block in response.content:
                            if hasattr(block, "text"):
                                print(f"\n{block.text}")
                        break

                    if response.stop_reason == "tool_use":
                        messages.append({"role": "assistant", "content": response.content})
                        tool_results = []
                        for block in response.content:
                            if block.type == "tool_use":
                                result = run_tool(block.name, block.input, notes, approved_competitors)
                                tool_results.append({
                                    "type": "tool_result",
                                    "tool_use_id": block.id,
                                    "content": str(result)
                                })
                        messages.append({"role": "user", "content": tool_results})

                # Recompute quality after extra research, show Gate 2 again
                final_quality = check_quality(notes, aspects)

    return notes


# --- MAIN ---

if __name__ == "__main__":
    company, aspects = get_user_inputs()
    run_agent(company, aspects)
